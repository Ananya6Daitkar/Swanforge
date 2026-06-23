from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT=Path(__file__).resolve().parents[1]
OUT=ROOT/"submission"/"BLACK_SWAN_FORGE_Moonshot_Paper.docx"
ACCENT=RGBColor(255,95,73); NAVY=RGBColor(12,23,36); BLUE=RGBColor(46,116,181); MUTED=RGBColor(92,107,128); WHITE=RGBColor(255,255,255)

def font(run,name="Aptos",size=11,bold=None,color=None,italic=None):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),name);run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"),name);run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color: run.font.color.rgb=color
    if italic is not None: run.italic=italic

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr();shd=tcPr.find(qn("w:shd")) or OxmlElement("w:shd");shd.set(qn("w:fill"),fill);tcPr.append(shd) if shd.getparent() is None else None

def margins(cell,top=100,start=140,bottom=100,end=140):
    tc=cell._tc.get_or_add_tcPr();tcMar=tc.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if tcMar.getparent() is None: tc.append(tcMar)
    for side,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tcMar.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}");node.set(qn("w:w"),str(val));node.set(qn("w:type"),"dxa");
        if node.getparent() is None:tcMar.append(node)

def set_repeat_header(row):
    trPr=row._tr.get_or_add_trPr();tblHeader=OxmlElement("w:tblHeader");tblHeader.set(qn("w:val"),"true");trPr.append(tblHeader)

def set_cell_width(cell,dxa):
    tcPr=cell._tc.get_or_add_tcPr();tcW=tcPr.find(qn("w:tcW")) or OxmlElement("w:tcW");tcW.set(qn("w:w"),str(dxa));tcW.set(qn("w:type"),"dxa");
    if tcW.getparent() is None:tcPr.append(tcW)

def table_geometry(table,widths):
    table.autofit=False;tblPr=table._tbl.tblPr;tblW=tblPr.find(qn("w:tblW"));tblW.set(qn("w:w"),str(sum(widths)));tblW.set(qn("w:type"),"dxa")
    ind=OxmlElement("w:tblInd");ind.set(qn("w:w"),"120");ind.set(qn("w:type"),"dxa");tblPr.append(ind)
    grid=table._tbl.tblGrid
    for c in list(grid):grid.remove(c)
    for w in widths:
        gc=OxmlElement("w:gridCol");gc.set(qn("w:w"),str(w));grid.append(gc)
    for row in table.rows:
        for i,cell in enumerate(row.cells):set_cell_width(cell,widths[i]);margins(cell);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def page_num(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT;run=paragraph.add_run();begin=OxmlElement("w:fldChar");begin.set(qn("w:fldCharType"),"begin");instr=OxmlElement("w:instrText");instr.set(qn("xml:space"),"preserve");instr.text=" PAGE ";end=OxmlElement("w:fldChar");end.set(qn("w:fldCharType"),"end");run._r.extend([begin,instr,end]);font(run,size=9,color=MUTED)

def add_p(doc,text="",style=None,bold_lead=None):
    p=doc.add_paragraph(style=style);p.paragraph_format.space_after=Pt(8);p.paragraph_format.line_spacing=1.25
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead);font(r,bold=True);r=p.add_run(text[len(bold_lead):]);font(r)
    else:r=p.add_run(text);font(r)
    return p

def add_callout(doc,label,text):
    t=doc.add_table(rows=1,cols=1);table_geometry(t,[9360]);c=t.cell(0,0);shade(c,"FFF2EF");p=c.paragraphs[0];r=p.add_run(label.upper()+"  ");font(r,size=9,bold=True,color=ACCENT);r=p.add_run(text);font(r,size=11,bold=True,color=NAVY);doc.add_paragraph().paragraph_format.space_after=Pt(2)

def add_h(doc,text,level=1):
    p=doc.add_paragraph(style=f"Heading {level}");p.add_run(text);return p

doc=Document();sec=doc.sections[0];sec.page_width=Inches(8.5);sec.page_height=Inches(11);sec.top_margin=Inches(1);sec.bottom_margin=Inches(1);sec.left_margin=Inches(1);sec.right_margin=Inches(1);sec.header_distance=Inches(.492);sec.footer_distance=Inches(.492)
styles=doc.styles
normal=styles["Normal"];normal.font.name="Aptos";normal.font.size=Pt(11);normal.paragraph_format.space_after=Pt(8);normal.paragraph_format.line_spacing=1.25
for name,size,before,after,color in [("Title",30,0,10,NAVY),("Subtitle",14,0,12,MUTED),("Heading 1",16,18,10,ACCENT),("Heading 2",13,12,6,BLUE),("Heading 3",12,8,4,NAVY)]:
    s=styles[name];s.font.name="Aptos Display" if name in ("Title","Heading 1") else "Aptos";s.font.size=Pt(size);s.font.bold=name!="Subtitle";s.font.color.rgb=color;s.paragraph_format.space_before=Pt(before);s.paragraph_format.space_after=Pt(after);s.paragraph_format.keep_with_next=True
title_ppr=styles["Title"]._element.get_or_add_pPr();title_border=title_ppr.find(qn("w:pBdr"));
if title_border is not None:title_ppr.remove(title_border)
for sname in ["List Bullet","List Number"]:
    s=styles[sname];s.font.name="Aptos";s.font.size=Pt(11);s.paragraph_format.left_indent=Inches(.375);s.paragraph_format.first_line_indent=Inches(-.194);s.paragraph_format.space_after=Pt(4);s.paragraph_format.line_spacing=1.208
header=sec.header.paragraphs[0];header.text="BLACK SWAN FORGE  /  MOONSHOT BLUEPRINT";font(header.runs[0],size=8,bold=True,color=MUTED)
page_num(sec.footer.paragraphs[0])

# Editorial cover
p=doc.add_paragraph();p.paragraph_format.space_before=Pt(115);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("MOONSHOT PAPER");font(r,size=10,bold=True,color=ACCENT)
p=doc.add_paragraph(style="Title");p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run("BLACK SWAN FORGE")
p=doc.add_paragraph(style="Subtitle");p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run("An adversarial world model for discovering unknown infrastructure collapses before civilization pays the cost")
add_callout(doc,"Thesis","Cities should be crash-tested against failures nobody has thought to rehearse.")
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(70);r=p.add_run("Synthetic research prototype  |  June 2026");font(r,size=10,bold=True,color=MUTED)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("Plausible under the assumptions of this model.");font(r,size=9,italic=True,color=MUTED)
doc.add_page_break()

add_h(doc,"Abstract")
add_p(doc,"Modern cities are not collections of independent utilities. They are tightly coupled systems in which authentication, payments, fuel, roads, hospitals, power, water, data, and emergency response can fail through delayed cross-domain dependencies. Yet most resilience practice begins with failures humans already know to test. Monitoring recognizes known deviations; digital twins replay expected behavior; chaos engineering injects selected faults. None is designed primarily to search for minimal, previously unimagined collapse mechanisms and then discover the cheapest intervention that breaks them.")
add_p(doc,"BLACK SWAN FORGE proposes Adversarial Infrastructure Science: an executable method that treats resilience as counterexample discovery. A deterministic city model propagates shocks through delayed conditional dependencies. Evolutionary search explores failure hypotheses. Delta debugging removes unnecessary conditions. Event-level parentage produces falsifiable causal traces. Cost-constrained intervention search replays counterfactual futures, while uncertainty analysis distinguishes robust mechanisms from fragile artifacts.")
add_p(doc,"The working prototype models 28 synthetic components and 46 dependencies across eleven domains. In a 20-seed equal-budget study, evolutionary search produced a mean 1.90x collapse-discovery ratio relative to uniform random sampling, with a normal-approximation 95% confidence interval of [1.65, 2.15], winning 19 of 20 seeds. A public IEEE 14-bus bridge adds DC power-flow redistribution and N-1 overload cascades. These results establish behavior inside the model; they are not claims about real-city likelihood.")
add_callout(doc,"Category","Discover unknown cascade -> minimize cause -> falsify links -> optimize repair -> verify survival.")

add_h(doc,"1. The problem humanity has misunderstood")
add_p(doc,"Civilization usually learns its most important infrastructure dependencies during emergencies. A hospital may possess generators yet depend on fuel deliveries; fuel stations may possess inventory yet depend on digital payments; payments may depend on telecom authentication; emergency response may depend on road access, dispatch data, and cellular coverage simultaneously. Each subsystem can appear locally resilient while the combined city is globally brittle.")
add_p(doc,"The mistaken assumption is that resilience is primarily an observation problem: collect more telemetry, produce better dashboards, and react faster. Observation is necessary, but it cannot warn us about a failure mechanism nobody has formulated. The deeper problem is epistemic. We do not know which cross-domain counterexamples we have failed to imagine.")
add_callout(doc,"First principle","A small initiating failure can be more dangerous than an obvious catastrophe because redundancy hides its consequences long enough for the cascade to cross organizational boundaries.")

add_h(doc,"2. Why existing approaches are insufficient")
table=doc.add_table(rows=1,cols=4);table.alignment=WD_TABLE_ALIGNMENT.LEFT;hdr=table.rows[0].cells
for i,t in enumerate(["Approach","Primary question","Strength","Missing capability"]):hdr[i].text=t;shade(hdr[i],"0C1724");font(hdr[i].paragraphs[0].runs[0],size=9,bold=True,color=WHITE)
rows=[("Monitoring","What is failing now?","Operational awareness","Cannot search unknown mechanisms"),("Digital twins","What will this known scenario do?","State and physics replay","Usually not adversarial or minimal"),("Chaos engineering","Does the system survive this chosen fault?","Empirical resilience testing","Starts from a human-selected hypothesis"),("BLACK SWAN FORGE","What failure nobody rehearsed can collapse the system?","Discovery, minimization and repair","Requires validated models before deployment")]
for row in rows:
    cells=table.add_row().cells
    for i,val in enumerate(row):cells[i].text=val;font(cells[i].paragraphs[0].runs[0],size=9);shade(cells[i],"FFF4F2" if row[0]=="BLACK SWAN FORGE" else "FFFFFF")
table_geometry(table,[1650,2050,2200,3460]);set_repeat_header(table.rows[0])

add_h(doc,"3. First-principles insight")
add_p(doc,"If a city is a dynamic directed system, then a disaster scenario is not a story. It is an executable counterexample: a bounded initial condition, a sequence of threshold crossings, a measurable critical outcome, and a counterfactual intervention. This changes the objective from predicting one future to searching the space of possible mechanisms.")
add_p(doc,"The method follows four principles:")
for text in ["Search rather than wait: generate candidate failures before reality does.","Minimize rather than dramatize: remove every condition not required for collapse.","Falsify rather than narrate: delete claimed causal links and replay the model.","Repair rather than score: identify the minimum-cost intervention that changes the outcome."]:
    p=doc.add_paragraph(style="List Number");r=p.add_run(text);font(r)

add_h(doc,"4. System architecture")
add_p(doc,"The prototype is a single deterministic TypeScript system. The simulation engine is independent of React and can be executed in tests or batch experiments. The visual interface consumes immutable snapshots rather than generating claims.")
for title,text in [("World model","Nodes represent infrastructure components with capacity, thresholds, recovery, criticality, zones, finite backups, status, failure reason, and time of failure. Directed edges encode dependency type, source requirement, impact strength, propagation delay, recovery condition, and explanation."),("Propagation engine","Each tick applies active shocks, executes due dependency effects, evaluates new threshold crossings, activates finite backups, updates recovery, records causal parents, and computes city metrics."),("Adversarial search","A seeded evolutionary algorithm mutates shock target, timing, duration, and intensity. Fitness rewards severity, cross-domain spread, depth, delay, and novelty while explicit constraints reject arbitrary destruction."),("Counterexample minimization","Delta debugging removes shocks and reduces intensity or duration while replaying the critical-outcome predicate."),("Intervention optimization","Single and combinatorial repairs are replayed. Non-dominated solutions form a cost-residual-risk frontier under deterministic and uncertain parameters.")]:
    add_h(doc,title,2);add_p(doc,text)

add_h(doc,"5. Formal model")
add_p(doc,"For component i, normalized capacity x_i(t) lies in [0,100]. Its discrete transition is the clipped sum of prior capacity, direct shocks, delayed parent effects, finite backup contribution, and recovery:")
add_callout(doc,"State transition","x_i(t+1) = clip[x_i(t) - s_i(t) - sum I_ji(t-d_ji) + b_i(t) + r_i(t), 0, 100]")
add_p(doc,"A dependency j->i fires when source capacity x_j falls below requirement theta_ji. The global collapse predicate becomes true when any hospital crosses its minimum threshold, water availability falls below 35%, or emergency response falls below 40%. Search maximizes severity, domains crossed, causal depth, delayed impact, and novelty minus initial-shock complexity and invalidity penalties.")
add_p(doc,"The minimizer seeks the smallest condition set that preserves collapse. The optimizer seeks the minimum intervention cost subject to the negation of that collapse predicate. A causal edge is necessary for the selected global outcome only if deleting it and replaying makes the predicate false.")

add_h(doc,"6. Demonstration: the quiet failure")
add_p(doc,"The deterministic demonstration begins with a telecom authentication outage. The immediate visual impact is limited. Two ticks later, payment authorization degrades. Fuel transactions then stop. Hospital reserves hide the downstream consequence until generator and ambulance fuel becomes unavailable. At T+9, hospital emergency capacity crosses its safety threshold.")
add_p(doc,"The system does not merely describe this chain. Every step references a simulation event and causal-parent identifier. Delta debugging reduces the initiating condition. Edge deletion shows that authentication-to-payments is necessary for the global outcome. The optimizer selects offline emergency fuel authorization at cost 2; counterfactual replay keeps the hospital operational.")
img=ROOT/"test-results"/"demo-complete.png"
if img.exists():
    doc.add_picture(str(img),width=Inches(6.45));cap=doc.add_paragraph("Figure 1. Working command-center replay after intervention discovery.");cap.alignment=WD_ALIGN_PARAGRAPH.CENTER;font(cap.runs[0],size=8,italic=True,color=MUTED)

add_h(doc,"7. Evaluation")
add_h(doc,"7.1 Equal-budget baseline",2)
table=doc.add_table(rows=1,cols=5);headers=["Method","Collapse cases","Unique signatures","Best severity","Median depth"]
for i,t in enumerate(headers):table.cell(0,i).text=t;shade(table.cell(0,i),"0C1724");font(table.cell(0,i).paragraphs[0].runs[0],size=8,bold=True,color=WHITE)
for row in [("Evolutionary","61","14","22","3"),("Uniform random","43","21","22","2"),("Criticality ranked","49","18","22","2"),("Graph centrality","54","16","23","2")]:
    cells=table.add_row().cells
    for i,v in enumerate(row):cells[i].text=v;font(cells[i].paragraphs[0].runs[0],size=9);shade(cells[i],"FFF4F2" if row[0]=="Evolutionary" else "FFFFFF")
table_geometry(table,[2500,1650,1800,1700,1710]);set_repeat_header(table.rows[0])
add_p(doc,"At 120 simulations per method, evolutionary search produced 61 collapse cases versus 43 for random sampling, a 1.42x ratio. Random sampling produced more unique signatures, revealing a real exploration-exploitation tradeoff rather than universal dominance.")
add_h(doc,"7.2 Multi-seed evidence",2)
add_p(doc,"Across 20 independent seeds at 60 simulations per method per seed, evolutionary search averaged 31.2 collapse cases versus 17.1 for random. The mean discovery ratio was 1.90x with a normal-approximation 95% confidence interval of [1.65, 2.15]; evolutionary search won 19 of 20 seeds. These settings are reproducible with npm run study.")
add_h(doc,"7.3 Robustness and sensitivity",2)
add_p(doc,"The telecom cascade persisted in 44 of 64 parameter perturbations (68.75%) when shock intensity varied by +/-20%, dependency strength by +/-25%, and recovery rate by +/-25%. It is labeled SENSITIVE, not robust. Spearman attribution assigned 82.7% of normalized absolute sensitivity to dependency strength, 17.1% to shock intensity, and 0.2% to recovery rate. The small-sample negative intensity association is disclosed as a threshold interaction, not a physical generalization.")

add_h(doc,"8. Physical-grid bridge")
add_p(doc,"The synthetic city model is complemented by a public IEEE 14-bus DC power-flow bridge using MATPOWER bus demand, active generation, and branch reactance. The solver fixes the slack-bus angle, solves B' theta = P, computes branch flows, removes overloaded lines, and repeats after an N-1 outage.")
add_p(doc,"Under explicitly derived experimental thermal limits, the worst N-1 replay begins with line 1-2, triggers three additional overload outages, and islands approximately 259 MW over two iterations. MATPOWER does not provide branch MVA limits for this fixture; therefore the overload limits are assumptions. The calculation demonstrates a credible integration boundary, not a validated prediction.")

add_h(doc,"9. Scientific honesty and limitations")
add_p(doc,"BLACK SWAN FORGE is a synthetic research prototype. It does not predict real cities. Dependencies, cross-domain parameters, intervention costs, and population mappings are demonstrative. The DC solver omits reactive power, voltage magnitudes, losses, and transient stability. The confidence interval is a normal approximation rather than a preregistered inferential study. Expert review has not yet been completed and must not be fabricated.")
add_callout(doc,"Required interpretation","Every result is plausible under the assumptions of this model—not evidence that a generated scenario will occur in reality.")

add_h(doc,"10. Path from prototype to reality")
for title,text in [("Open benchmarks","Validate search on public power, water, transport, and communications test networks."),("Secure ingestion","Import operator-owned topology and parameter distributions inside the operator's security boundary."),("Hybrid solvers","Replace scalar capacities with calibrated power-flow, hydraulic, traffic, inventory, queueing, and clinical models."),("Expert-blinded evaluation","Compare evolutionary and baseline cascades on plausibility, novelty, causal coherence, actionability, and false-confidence risk."),("Shadow exercises","Run against historical and tabletop scenarios without operational control."),("Federated resilience network","Exchange privacy-preserving dependency contracts across operators while sensitive topology remains local.")]:
    add_h(doc,title,2);add_p(doc,text)

add_h(doc,"11. Long-term implication")
add_p(doc,"Today, cities discover hidden dependencies when emergencies expose them. If this method succeeds, infrastructure design changes from compliance against a finite checklist to continuous adversarial inquiry. Hospitals could test supplier, payment, road, telecom, and generator assumptions together. Utilities could exchange dependency contracts without exposing full topology. Regulators could require executable counterexamples and counterfactual repairs rather than static risk registers.")
add_p(doc,"The long-term ambition is a crash-testing layer for civilization: not an autonomous controller, but a machine for generating hypotheses that human experts can falsify before real people become the experiment. The future it attempts to create is one in which systemic fragility is discovered computationally, debated explicitly, and repaired deliberately.")
add_callout(doc,"Moonshot","Civilization should learn its hidden dependencies in simulation, not in disaster.")

doc.add_page_break();add_h(doc,"References")
refs=["MATPOWER. IEEE 14-bus test case, case14.m. https://github.com/MATPOWER/matpower/blob/master/data/case14.m","Zimmerman, R. D., Murillo-Sanchez, C. E., and Thomas, R. J. MATPOWER: Steady-State Operations, Planning, and Analysis Tools for Power Systems Research and Education.","Zeller, A., and Hildebrandt, R. Simplifying and Isolating Failure-Inducing Input. IEEE Transactions on Software Engineering, 2002.","BLACK SWAN FORGE repository documentation: Architecture, Evaluation, Formal Model, Limitations, and Expert Review Protocol."]
for ref in refs:
    p=doc.add_paragraph();p.paragraph_format.left_indent=Inches(.2);p.paragraph_format.first_line_indent=Inches(-.2);p.paragraph_format.space_after=Pt(9);r=p.add_run(ref);font(r,size=9)

doc.core_properties.title="BLACK SWAN FORGE - Moonshot Paper";doc.core_properties.subject="Adversarial Infrastructure Science";doc.core_properties.author="BLACK SWAN FORGE Team";doc.core_properties.keywords="infrastructure, adversarial search, simulation, causal minimization, resilience"
OUT.parent.mkdir(parents=True,exist_ok=True);doc.save(OUT);print(OUT)
